import shutil
from docx import Document

class Student:
    """docstring for Student"""
    def __init__(self, args):
        self.id = args[0]
        self.name = args[1]
        self.subject = args[2]
        self.thesis = args[3]
        self.supvisor = args[4]

    def introduce(self):
        return f"My name is {self.name} and my supervisor is {self.supvisor}."


def copy_file(source_path, destination_path):
    shutil.copy(source_path, destination_path)

def replace_text_in_docx(docx_file, old_text, new_text):
    doc = Document(docx_file)
    # for paragraph in doc.paragraphs:
    #     if old_text in paragraph.text:
    #         paragraph.text = paragraph.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    # 记录原始文本的格式
                    original_format = cell.paragraphs[0].alignment
                    cell.text = cell.text.replace(old_text, new_text)
                    cell.paragraphs[0].alignment = original_format
    doc.save(docx_file)

def parse_data(file_path):
    with open(file_path, 'r') as file:
        for line in file:
            # 使用 split() 方法根据制表符分隔每一行，并存入数组
            split_line = line.strip().split('\t')
            # print(split_line)
            student = Student(split_line)
            # print(student.introduce())
            filename = student.id + student.name + ".docx"
            copy_file("./temp.docx", filename)
            replace_text_in_docx(filename, "11111", student.name)
            replace_text_in_docx(filename, "22222", student.id)
            replace_text_in_docx(filename, "33333", student.supvisor)
            replace_text_in_docx(filename, "44444", student.subject)
            replace_text_in_docx(filename, "55555", student.thesis)


def main():
    file_path = 'data.txt'  # 请将文件路径替换为实际的文件路径
    lines = parse_data(file_path)
    # print(lines)

if __name__ == "__main__":
    main()
